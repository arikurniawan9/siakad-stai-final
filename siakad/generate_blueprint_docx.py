import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=160, right=160):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def build_docx():
    doc = docx.Document()
    
    # Page setup - A4 with 1 inch margin
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        
        # Header & Footer
        header = section.header
        p_head = header.paragraphs[0]
        p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h_run = p_head.add_run("STAI AL-ITTIHAD CIANJUR — CETAK BIRU & PEMBIAYAAN SISTEM SIAKAD & LMS")
        h_run.font.name = "Calibri"
        h_run.font.size = Pt(8.5)
        h_run.font.color.rgb = RGBColor(140, 150, 160)

        footer = section.footer
        p_foot = footer.paragraphs[0]
        p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = p_foot.add_run("Dokumen Resmi Spesifikasi Teknis & Rencana Anggaran Biaya (RAB) — 2026")
        f_run.font.name = "Calibri"
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGBColor(140, 150, 160)

    # Color Palette - Islamic Enterprise (Deep Navy, Emerald Green, Slate Gray)
    COLOR_PRIMARY = RGBColor(27, 54, 93)     # Deep Navy #1B365D
    COLOR_SECONDARY = RGBColor(16, 124, 65)  # Islamic Emerald #107C41
    COLOR_DARK = RGBColor(30, 41, 59)        # Slate Dark #1E293B
    COLOR_MUTED = RGBColor(100, 116, 139)    # Muted Slate #64748B
    
    HEX_HEADER = "1B365D"
    HEX_SUBHEADER = "107C41"
    HEX_LIGHT_ROW = "F8FAFC"
    HEX_ACCENT_ROW = "F1F5F9"
    HEX_CALLOUT_BG = "ECFDF5"

    # ==================== COVER / JUDUL UTAMA ====================
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_inst = title_p.add_run("PROPOSAL BLUEPRINT ARSITEKTUR & RENCANA PEMBIAYAAN\n")
    r_inst.font.name = "Arial"
    r_inst.font.size = Pt(13)
    r_inst.font.bold = True
    r_inst.font.color.rgb = COLOR_SECONDARY
    
    r_title = title_p.add_run("SISTEM INFORMASI AKADEMIK (SIAKAD) &\nLEARNING MANAGEMENT SYSTEM (SALAM LMS)")
    r_title.font.name = "Arial Black"
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(8)
    sub_p.paragraph_format.space_after = Pt(20)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Integrasi Tata Kelola Kampus, Billing Virtual Account Bank Syariah Indonesia (BSI), & Platform Pembelajaran Terpadu")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = COLOR_MUTED

    # Meta Table Box
    meta_tbl = doc.add_table(rows=6, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False
    
    col_widths = [Inches(2.3), Inches(3.9)]
    meta_data = [
        ("Institusi Klien", "Sekolah Tinggi Agama Islam (STAI) Al-Ittihad Cianjur"),
        ("Lingkup Sistem", "SIAKAD Enterprise + SALAM LMS (Sinkronisasi Dua Arah)"),
        ("Teknologi Resmi", "Laravel + Inertia.js (React) + PostgreSQL 16 + Tailwind CSS"),
        ("Integrasi Finansial", "Host-to-Host Virtual Account (VA) Bank Syariah Indonesia (BSI)"),
        ("Total Nilai Pembiayaan", "Rp 25.000.000,- (Dua Puluh Lima Juta Rupiah)"),
        ("Tahun Pelaksanaan", "2026 (Production Deployment & Training)")
    ]
    
    for row_idx, (label, val) in enumerate(meta_data):
        row = meta_tbl.rows[row_idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = col_widths[0]
        cell_val.width = col_widths[1]
        
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_before = Pt(3)
        p_lbl.paragraph_format.space_after = Pt(3)
        r1 = p_lbl.add_run(label)
        r1.font.name = "Calibri"
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_PRIMARY
        
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(3)
        p_val.paragraph_format.space_after = Pt(3)
        r2 = p_val.add_run(val)
        r2.font.name = "Calibri"
        r2.font.size = Pt(9.5)
        if "25.000.000" in val:
            r2.font.bold = True
            r2.font.color.rgb = COLOR_SECONDARY
        elif "Laravel" in val:
            r2.font.bold = True
            r2.font.color.rgb = COLOR_PRIMARY
            
        set_cell_background(cell_lbl, "F1F5F9")
        set_cell_background(cell_val, "FFFFFF")
        set_cell_margins(cell_lbl, 90, 90, 110, 110)
        set_cell_margins(cell_val, 90, 90, 110, 110)
        
    set_table_borders(meta_tbl, "CBD5E1")
    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Helper formatters
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(13.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_PRIMARY
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = COLOR_SECONDARY
        return h

    def add_paragraph(text, bold_prefix=None, space_after=5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = "Calibri"
            r_bold.font.size = Pt(10)
            r_bold.font.bold = True
            r_bold.font.color.rgb = COLOR_DARK
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_DARK
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = "Calibri"
            r_bold.font.size = Pt(10)
            r_bold.font.bold = True
            r_bold.font.color.rgb = COLOR_DARK
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_DARK
        return p

    # ==================== BAB 1: RINGKASAN EKSEKUTIF ====================
    add_heading_1("1. RINGKASAN EKSEKUTIF & LATAR BELAKANG")
    add_paragraph("Dalam rangka mewujudkan digitalisasi tata kelola perguruan tinggi berbasis syariah yang modern, transparan, dan akuntabel di STAI Al-Ittihad Cianjur, diperlukan ekosistem piranti lunak terpadu yang menghubungkan administrasi akademik (SIAKAD), gerbang pembayaran perbankan syariah (Bank Syariah Indonesia / BSI), serta platform pembelajaran interaktif harian (SALAM LMS).")
    add_paragraph("Proyek ini mencakup pengembangan lengkap dari sisi basis data relasional PostgreSQL, antarmuka pengguna responsif (UI/UX) berbasis React, backend Laravel teroptimasi, gerbang Host-to-Host Virtual Account BSI, serta modul sinkronisasi otomatis dua arah hingga aplikasi siap digunakan secara penuh (Turnkey Production Solution).")

    # ==================== BAB 2: ARSITEKTUR & TECH STACK ====================
    add_heading_1("2. ARSITEKTUR SISTEM & SPESIFIKASI TECH STACK RESMI")
    add_paragraph("Untuk memastikan aplikasi memiliki performa tinggi, sangat hemat penggunaan RAM di VPS, bebas lelet saat serbuan pengisian KRS, dan mudah dirawat, disepakati penggunaan arsitektur standar industri:")

    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tech_table.autofit = False
    t_widths = [Inches(1.8), Inches(2.2), Inches(2.2)]
    
    t_hdr = tech_table.rows[0].cells
    for i, title in enumerate(["Layer Sistem", "Teknologi / Framework", "Peran & Keunggulan Utama"]):
        t_hdr[i].width = t_widths[i]
        p = t_hdr[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(t_hdr[i], HEX_HEADER)
        set_cell_margins(t_hdr[i], 80, 80, 100, 100)

    tech_items = [
        ("Backend Core Engine", "Laravel (PHP 8.3/8.4)", "Routing, Eloquent ORM, DB Migrations, Background Queue, & Task Scheduler."),
        ("Frontend Bridge", "Inertia.js v2", "Menghubungkan Laravel & React mulus tanpa ribet menulis API boilerplate."),
        ("Frontend UI Client", "React + Tailwind CSS", "Antarmuka Single Page Application yang interaktif, cepat, dan serasi dg LMS."),
        ("Database Relasional", "PostgreSQL 16", "Database tangguh dengan transaksi ACID, integritas foreign key, & JSONB."),
        ("Virtual Account Gateway", "BSI Open API (H2H)", "Inquiry & Payment Callback Webhook real-time dengan verifikasi signature."),
        ("Penyimpanan Berkas", "MinIO S3 / Storage", "Penyimpanan dokumen PMB, ijazah, materi, & berkas tugas mahasiswa."),
        ("Web Server & Proxy", "Nginx + PHP-FPM", "Reverse proxy berkinerja tinggi, SSL HTTPS Let's Encrypt, & OPcache (~80MB RAM).")
    ]

    for item in tech_items:
        row = tech_table.add_row()
        for c_idx, val in enumerate(item):
            cell = row.cells[c_idx]
            cell.width = t_widths[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            if c_idx == 0:
                r.font.bold = True
            set_cell_background(cell, HEX_LIGHT_ROW if tech_items.index(item) % 2 == 1 else "FFFFFF")
            set_cell_margins(cell, 70, 70, 90, 90)

    set_table_borders(tech_table, "CBD5E1")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ==================== BAB 3: MODUL UTAMA SIAKAD ====================
    add_heading_1("3. MATRIKS MODUL & FITUR UTAMA SIAKAD TERPADU")
    
    add_heading_2("3.1 Master Infrastruktur: Gedung & Ruang Kelas")
    add_bullet(" Manajemen denah gedung, jumlah lantai, lokasi kampus, dan pengelompokan fasilitas.", "Master Gedung:")
    add_bullet(" Kode ruang, nama ruang, kapasitas kursi kuliah reguler vs ujian, tipe ruang (Teori, Lab Komputer, Microteaching, Auditorium), dan inventaris fasilitas (AC, Proyektor, Sound, CCTV).", "Master Ruang Kelas:")
    add_bullet(" Sistem memvalidasi dan mencegah otomatis bentrok ruang dan jam perkuliahan pada saat penyusunan jadwal.", "Anti-Clash Scheduler:")

    add_heading_2("3.2 Master Akademik, Kurikulum & Tugas Struktural")
    add_bullet(" Pengaturan kalender akademik, semester (Ganjil, Genap, Pendek), masa bayar SPP, masa KRS, masa revisi KPRS, masa input nilai, dan masa EDOM.", "Tahun & Periode Akademik:")
    add_bullet(" S1 Pendidikan Agama Islam (PAI), S1 Manajemen Pendidikan Islam (MPI), S1 Hukum Ekonomi Syariah (HES), S1 Pendidikan Guru MI (PGMI), S1 Ekonomi Syariah (ESY) beserta struktur kurikulum dan bobot SKS.", "Program Studi & Kurikulum:")
    add_bullet(" Pendataan Ketua STAI, Para Wakil Ketua, Kaprodi, Sekretaris Prodi, Kepala BAAK, Kepala Bagian Keuangan, dan Dosen PA lengkap dengan Nomor SK dan otoritas tanda tangan digital.", "Tugas & Jabatan Struktural:")

    add_heading_2("3.3 PMB Online & Otomasi Virtual Account Billing BSI")
    add_bullet(" Formulir pendaftaran mandiri calon mahasiswa, pilihan jalur (Reguler, Tahfidz, Prestasi, Pindahan), dan pemilihan prodi 1 & 2.", "Pendaftaran Calon Mahasiswa:")
    add_bullet(" Begitu calon mahasiswa submit formulir pendaftaran, sistem secara otomatis menerbitkan tagihan PMB dan nomor Virtual Account BSI resmi (Prefix 9928 + Kode 01 + ID Registrasi, contoh: 992801260001).", "Auto-Generate VA BSI PMB:")
    add_bullet(" Begitu pembayaran diterima melalui BSI Mobile/ATM/Teller, Webhook BSI langsung mengupdate status invoice menjadi LUNAS dan otomatis membuka akses unggah berkas persyaratan (Ijazah, KTP, KK, SKCK, Foto).", "Webhook Callback Real-Time:")
    add_bullet(" Calon mahasiswa yang lulus seleksi otomatis diterbitkan NIM resmi dan akun portal akademik mahasiswa.", "Penerbitan NIM & Akun:")

    add_heading_2("3.4 Modul Keuangan & Billing Tagihan Terpadu (SPP / UKT)")
    add_bullet(" Staf Keuangan dapat men-generate tagihan SPP/UKT satu semester penuh untuk seluruh mahasiswa aktif hanya dengan 1-klik.", "Mass Billing Engine:")
    add_bullet(" Format VA UKT: 9928 + 02 + NIM (contoh: 99280221010042). Dilengkapi simulasi sandbox untuk testing lokal.", "Format VA BSI Mahasiswa:")
    add_bullet(" Mahasiswa yang belum lunas UKT semester berjalan secara otomatis terkunci dari akses pengisian KRS, cetak kartu ujian, dan cetak KHS.", "Financial Lock Guard:")
    add_bullet(" Modul pengajuan keringanan/cicilan atas persetujuan pimpinan untuk membuka kunci akademik sementara.", "Dispensasi Keuangan:")

    add_heading_2("3.5 Modul KRS Online & Dosen Pembimbing Akademik (Dosen PA)")
    add_bullet(" Batas SKS dihitung otomatis dari IPS semester sebelumnya (IPS >= 3.50 maks 24 SKS, IPS 3.00-3.49 maks 22 SKS, dst.).", "Validasi Beban SKS:")
    add_bullet(" Dashboard dosen wali untuk memantau progres studi mahasiswa bimbingan, menyetujui (Approve), atau meminta revisi KRS.", "Portal Dosen PA:")
    add_bullet(" Menghasilkan dokumen PDF KRS resmi ber-kop institusi STAI Al-Ittihad dan barcode validasi.", "Cetak Lembar KRS:")

    add_heading_2("3.6 Modul EDOM (Evaluasi Dosen Oleh Mahasiswa)")
    add_bullet(" Kuesioner evaluasi 4 kompetensi dosen (Pedagogik, Profesional, Kepribadian, Sosial) dengan skala 1-5.", "Instrumen EDOM:")
    add_bullet(" Identitas pengisi dirahasiakan total (tidak mencatat student_id pada lembar jawaban) untuk menjamin objektivitas evaluasi.", "Prinsip 100% Anonim:")
    add_bullet(" Pengisian EDOM menjadi syarat mutlak membuka KHS atau mengisi KRS semester selanjutnya.", "EDOM Lock Guard:")
    add_bullet(" Analitik skor indeks kinerja dosen untuk Kaprodi, Ketua STAI, dan borang akreditasi.", "Laporan Mutu Dosen:")

    add_heading_2("3.7 Modul KHS, Transkrip Akademik, & Yudisium")
    add_bullet(" Input nilai komponen (Presensi, Tugas, UTS, UAS) atau ditarik otomatis dari SALAM LMS.", "Buku Nilai & KHS:")
    add_bullet(" Cetak KHS dan Transkrip Lengkap ber-kop resmi dan Dynamic QR Code Verification Portal.", "Transkrip Digital:")
    add_bullet(" Verifikasi otomatis pemenuhan minimal 144 SKS, bebas nilai E, dan tes Tahfidz.", "Validasi Yudisium:")

    add_heading_2("3.8 Superadmin Master Control & Mode Menyamar (Role Impersonation)")
    add_bullet(" Superadmin dapat menyamar dan melihat antarmuka sebagai Mahasiswa, Dosen, Kaprodi, atau Staf Keuangan secara instan tanpa perlu mengetahui password target.", "Mode Menyamar (Impersonate):")
    add_bullet(" Banner kuning emas aktif di bagian atas layar dengan tombol instan 'Kembali ke Akun Superadmin', seluruh aksi dicatat di Audit Log.", "Keamanan & Audit Penyamaran:")
    add_bullet(" Backup database PostgreSQL on-demand (.sql.gz), monitor latency database pool, dan pengelolaan kapasitas MinIO S3.", "Database & Storage Control:")
    add_bullet(" Mengunci akses publik saat pembaruan sistem dengan whitelist IP admin.", "Maintenance Mode:")

    add_heading_2("3.9 Sistem Autentikasi & Captcha 4-Digit")
    add_bullet(" Mendukung NIM, NIDN, NIP, Username, dan Email resmi kampus.", "Multi-Identifier Login:")
    add_bullet(" Kode alfanumerik 4-digit acak bebas karakter ambigu (23456789ABCDEFGHJKLMNPQRSTUVWXYZ).", "Engine Captcha 4-Digit:")
    add_bullet(" Input otomatis berubah kapital (uppercase) dan diverifikasi melalui encrypted session cookie sekali pakai (single-use).", "Session & Auto-Uppercase:")

    # ==================== BAB 4: SALAM LMS & SINKRONISASI ====================
    add_heading_1("4. PLATFORM PEMBELAJARAN SALAM LMS & SINKRONISASI DUA ARAH")
    add_paragraph("Platform SALAM LMS yang telah disiapkan mencakup fitur unggulan:")
    add_bullet(" Kode QR berputar per 20 detik anti titip absen + Passcode 6-digit darurat + Mode Proyektor Kelas.", "Presensi Dynamic QR:")
    add_bullet(" Gerbang integritas, auto-fullscreen guard, deteksi pindah tab (maks 3x toleransi -> auto force-submit), blokir shortcut F12/Inspect/Copy-Paste.", "CBT Anti-Cheating Lockdown:")
    add_bullet(" Unggah berkas ke Object Storage MinIO S3 + Penilaian Rubrik OBE 4 Preset Resmi STAI Al-Ittihad.", "Tugas & Rubrik OBE:")
    add_bullet(" SIAKAD Laravel mem-push master prodi, kurikulum, dosen, kelas, dan mahasiswa aktif ke LMS. Sebaliknya, LMS mem-push rekap nilai gradebook dan persentase presensi ke SIAKAD.", "Engine Sinkronisasi Dua Arah:")

    # ==================== BAB 5: RENCANA ANGGARAN BIAYA ====================
    add_heading_1("5. RENCANA ANGGARAN BIAYA (RAB) & PEMBIAYAAN PROYEK")
    add_paragraph("Total nilai investasi pengembangan perangkat lunak terpadu SIAKAD (Laravel + Inertia React) & SALAM LMS STAI Al-Ittihad Cianjur adalah sebesar:")
    
    # Big Callout Total
    p_tot = doc.add_paragraph()
    p_tot.paragraph_format.space_before = Pt(6)
    p_tot.paragraph_format.space_after = Pt(12)
    p_tot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_box = p_tot.add_run("TOTAL NILAI KONTRAK: Rp 25.000.000,-\n(Dua Puluh Lima Juta Rupiah)")
    r_box.font.name = "Arial Black"
    r_box.font.size = Pt(14)
    r_box.font.bold = True
    r_box.font.color.rgb = COLOR_SECONDARY

    add_paragraph("Berikut adalah rincian alokasi biaya per modul dan komponen pekerjaan:")

    rab_table = doc.add_table(rows=1, cols=4)
    rab_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rab_table.autofit = False
    rab_widths = [Inches(0.6), Inches(3.2), Inches(1.3), Inches(1.3)]
    
    hdr_cells = rab_table.rows[0].cells
    headers = ["No", "Komponen Pekerjaan & Modul Sistem", "Alokasi (%)", "Biaya (IDR)"]
    for i, title in enumerate(headers):
        hdr_cells[i].width = rab_widths[i]
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], HEX_HEADER)
        set_cell_margins(hdr_cells[i], 100, 100, 100, 100)

    rab_items = [
        ("1", "Pondasi Arsitektur, Skema PostgreSQL 16, Laravel Core, & Inertia.js React", "10%", "Rp 2.500.000"),
        ("2", "Sistem Keamanan Login, Engine Captcha 4-Digit, & Superadmin Impersonation", "10%", "Rp 2.500.000"),
        ("3", "Modul Master Gedung, Ruang Kelas, Tahun Akademik, Kurikulum, & Tugas Struktural", "12%", "Rp 3.000.000"),
        ("4", "Modul PMB Online & Integrasi Otomasi VA Billing Bank Syariah Indonesia (BSI)", "16%", "Rp 4.000.000"),
        ("5", "Modul Keuangan SPP/UKT Massal, H2H BSI Webhook Callback, & Financial Lock", "14%", "Rp 3.500.000"),
        ("6", "Modul KRS Online, Dosen PA Approval, EDOM 4 Kompetensi, & KHS / Transkrip", "14%", "Rp 3.500.000"),
        ("7", "Modul Integrasi SALAM LMS, Presensi QR, CBT Anti-Cheat, & Dual-Way Sync", "10%", "Rp 2.500.000"),
        ("8", "Modul Superadmin Maintenance, Database Backup, & System Settings Panel", "6%", "Rp 1.500.000"),
        ("9", "Setup VPS Hosting, Docker Multi-Stage, Nginx Reverse Proxy, & SSL Certbot", "4%", "Rp 1.000.000"),
        ("10", "User Acceptance Testing (UAT), Pelatihan Administrator/Dosen, & Dokumentasi", "4%", "Rp 1.000.000"),
    ]

    for item in rab_items:
        row = rab_table.add_row()
        for c_idx, val in enumerate(item):
            cell = row.cells[c_idx]
            cell.width = rab_widths[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif c_idx in (2, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(9)
            set_cell_background(cell, HEX_LIGHT_ROW if int(item[0]) % 2 == 1 else "FFFFFF")
            set_cell_margins(cell, 80, 80, 100, 100)

    # Total Row
    tot_row = rab_table.add_row()
    for c_idx, val in enumerate(["", "TOTAL KESELURUHAN BIAYA PENGEMBANGAN", "100%", "Rp 25.000.000"]):
        cell = tot_row.cells[c_idx]
        cell.width = rab_widths[c_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        if c_idx in (2, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(val)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_PRIMARY
        set_cell_background(cell, HEX_ACCENT_ROW)
        set_cell_margins(cell, 100, 100, 100, 100)

    set_table_borders(rab_table, "CBD5E1")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ==================== BAB 6: TERMIN PEMBAYARAN ====================
    add_heading_1("6. TAHAPAN & TERMIN PEMBAYARAN (PAYMENT MILESTONES)")
    add_paragraph("Mekanisme pembayaran dibagi ke dalam 3 (tiga) termin berdasarkan pencapaian milestone pekerjaan:")

    termin_table = doc.add_table(rows=1, cols=4)
    termin_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    termin_table.autofit = False
    t_widths = [Inches(1.1), Inches(1.1), Inches(3.0), Inches(1.2)]
    
    t_hdr_cells = termin_table.rows[0].cells
    t_headers = ["Termin", "Persentase", "Uraian Milestone & Syarat Pencairan", "Nominal (IDR)"]
    for i, title in enumerate(t_headers):
        t_hdr_cells[i].width = t_widths[i]
        p = t_hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1) else (WD_ALIGN_PARAGRAPH.RIGHT if i == 3 else WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(title)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(t_hdr_cells[i], HEX_SUBHEADER)
        set_cell_margins(t_hdr_cells[i], 100, 100, 100, 100)

    termin_items = [
        ("Termin I (DP)", "30%", "Uang Muka Kerja saat kontrak ditandatangani, inisialisasi arsitektur, setup Laravel + Inertia React, migrasi PostgreSQL 16, Captcha 4-digit, dan Master Gedung/Ruang.", "Rp 7.500.000"),
        ("Termin II (Mid)", "40%", "Selesainya Modul PMB Online + Integrasi VA BSI, Keuangan SPP/UKT, Modul KRS Online + Dosen PA, serta Modul EDOM & Superadmin Control.", "Rp 10.000.000"),
        ("Termin III (Final)", "30%", "Selesainya integrasi sinkronisasi SALAM LMS, pengujian User Acceptance Testing (UAT), deployment ke VPS Hosting, pelatihan civitas, dan serah terima akhir.", "Rp 7.500.000"),
    ]

    for t_item in termin_items:
        row = termin_table.add_row()
        for c_idx, val in enumerate(t_item):
            cell = row.cells[c_idx]
            cell.width = t_widths[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if c_idx in (0, 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif c_idx == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx in (0, 3):
                r.font.bold = True
            set_cell_background(cell, "FFFFFF")
            set_cell_margins(cell, 80, 80, 100, 100)

    set_table_borders(termin_table, "CBD5E1")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ==================== BAB 7: JADWAL & SLA ====================
    add_heading_1("7. JADWAL PENGERJAAN & SERVICE LEVEL AGREEMENT (SLA)")
    add_paragraph("Target durasi pelaksanaan proyek adalah selama 4 (Empat) Pekan Kerja kalender dengan rincian linimasa sebagai berikut:")
    add_bullet(" Setup Laravel, Inertia.js React, Auth Captcha 4-digit, Master Gedung/Ruang, Tahun Akademik, dan Tugas Struktural.", "Pekan 1 (Core & Security):")
    add_bullet(" Portal PMB Calon Mahasiswa, Engine VA BSI, Webhook Callback BSI, Bulk Tagihan SPP, dan Financial Lock Guard.", "Pekan 2 (PMB & Fintech VA BSI):")
    add_bullet(" Formulir KRS Online, Portal Dosen PA Approval, Kuesioner EDOM 4 Kompetensi, Buku Nilai, dan KHS / Transkrip.", "Pekan 3 (Akademik, KRS & EDOM):")
    add_bullet(" Gateway Sinkronisasi Dua Arah ke SALAM LMS, Superadmin Maintenance/Backup, Dockerisasi, Deployment ke VPS Hosting, UAT, dan Pelatihan.", "Pekan 4 (LMS Sync & Deployment):")

    add_heading_2("Garansi & Masa Pemeliharaan (Maintenance SLA)")
    add_bullet(" Disediakan masa pemeliharaan dan perbaikan bug cuma-cuma (*free warranty*) selama **3 (Tiga) Bulan** sejak Berita Acara Serah Terima (BAST).", "Garansi Bebas Bug:")
    add_bullet(" Bantuan teknis respon cepat untuk kendala sistem operasional harian melalui tim pengembang.", "Technical Support:")
    add_bullet(" Penyerahan seluruh Source Code utuh, file blueprint arsitektur, spesifikasi API VA BSI, dan Buku Panduan Pengguna (User Manual Book).", "Kepemilikan Aset:")

    # ==================== SIGNATURE BOX ====================
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    sig_tbl = doc.add_table(rows=3, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_tbl.autofit = False
    sig_widths = [Inches(3.2), Inches(3.2)]
    
    r0 = sig_tbl.rows[0]
    p_s1 = r0.cells[0].paragraphs[0]
    p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s1 = p_s1.add_run("Pihak Pertama / Klien:\nSTAI Al-Ittihad Cianjur")
    r_s1.font.bold = True
    r_s1.font.size = Pt(9.5)
    
    p_s2 = r0.cells[1].paragraphs[0]
    p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s2 = p_s2.add_run("Pihak Kedua / Pengembang:\nTim Pengembang Software SIAKAD & LMS")
    r_s2.font.bold = True
    r_s2.font.size = Pt(9.5)

    r1 = sig_tbl.rows[1]
    r1.cells[0].paragraphs[0].text = "\n\n\n"
    r1.cells[1].paragraphs[0].text = "\n\n\n"

    r2 = sig_tbl.rows[2]
    p_e1 = r2.cells[0].paragraphs[0]
    p_e1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_e1 = p_e1.add_run("( _______________________ )\nKetua / Pimpinan STAI Al-Ittihad")
    r_e1.font.size = Pt(9)
    
    p_e2 = r2.cells[1].paragraphs[0]
    p_e2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_e2 = p_e2.add_run("( _______________________ )\nLead Software Architect / Developer")
    r_e2.font.size = Pt(9)

    for row in sig_tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = sig_widths[i]
            set_cell_background(cell, "FFFFFF")

    # Save document - save to multiple names so locked file doesn't block
    paths = [
        r"C:\project\salam-siakad\siakad\BLUEPRINT_DAN_RENCANA_PEMBIAYAAN_SIAKAD_LMS.docx",
        r"C:\project\salam-siakad\siakad\BLUEPRINT_DAN_RENCANA_PEMBIAYAAN_SIAKAD_LMS_TERBARU.docx"
    ]
    
    for path in paths:
        try:
            doc.save(path)
            print(f"SUCCESS: Saved to {path}")
        except PermissionError:
            print(f"NOTICE: {path} is open in Word, skipping.")

if __name__ == "__main__":
    build_docx()
